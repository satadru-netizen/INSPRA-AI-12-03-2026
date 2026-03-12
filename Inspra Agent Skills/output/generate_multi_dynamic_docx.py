"""
Generate two .docx files for the Multi Dynamic Oren Park project:
1. MultiDynamic_SystemPrompt_OrenPark.docx
2. MultiDynamic_DemoInstructions_OrenPark.docx
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def set_footer(doc, text="Prepared by Inspra AI"):
    """Add a centered gray footer to all sections."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.font.name = "Calibri"


def set_margins(doc, inches=1.0):
    """Set 1-inch margins on all sections."""
    for section in doc.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


def add_shaded_monospace_paragraph(doc, text):
    """Add a paragraph with Courier New font and light gray/blue shading (D9E2F3)."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    return p


def set_run_font(run, name="Calibri", size=11):
    run.font.name = name
    run.font.size = Pt(size)
    return run


# ---------------------------------------------------------------------------
# FILE 1: System Prompt
# ---------------------------------------------------------------------------
def generate_system_prompt_docx():
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    set_margins(doc)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(
        "Multi Dynamic Oren Park \u2014 Outbound Voice Agent System Prompt"
    )
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = "Calibri"

    doc.add_paragraph()  # spacer

    # Read the markdown file
    md_path = os.path.join(OUTPUT_DIR, "Multi Dynamic - System Prompt.md")
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    # Define the sections we want to extract
    section_names = [
        "Identity",
        "Context",
        "Conversation Principles",
        "Voice Format Rules",
        "Conversation Flow",
        "Objection Handling",
        "Tools Available",
        "Guardrails",
        "Boundaries",
    ]

    # Split content by ## headings
    sections = re.split(r"^## ", content, flags=re.MULTILINE)

    for section_text in sections:
        if not section_text.strip():
            continue

        # Extract heading (first line)
        lines = section_text.strip().split("\n", 1)
        heading = lines[0].strip()

        # Skip the top-level title and horizontal rules
        if heading.startswith("#") or heading == "---":
            continue

        # Check if this heading matches one of our sections
        matched = False
        for sn in section_names:
            if heading.lower() == sn.lower():
                matched = True
                break

        if not matched:
            continue

        # Add Heading 2
        h = doc.add_heading(heading, level=2)
        for run in h.runs:
            run.font.size = Pt(13)
            run.font.name = "Calibri"
            run.bold = True

        # Add body content in monospace with shading
        body = lines[1].strip() if len(lines) > 1 else ""
        if body:
            paragraphs = re.split(r"\n\n+", body)
            for para_text in paragraphs:
                para_text = para_text.strip()
                if para_text:
                    add_shaded_monospace_paragraph(doc, para_text)

    # Footer
    set_footer(doc)

    # Save
    out_path = os.path.join(OUTPUT_DIR, "MultiDynamic_SystemPrompt_OrenPark.docx")
    doc.save(out_path)
    print(f"Created: {out_path}")


# ---------------------------------------------------------------------------
# FILE 2: Demo Instructions
# ---------------------------------------------------------------------------
def generate_demo_instructions_docx():
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    set_margins(doc)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(
        "Multi Dynamic Oren Park \u2014 Prototype Demo Instructions"
    )
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = "Calibri"

    doc.add_paragraph()  # spacer

    # Metadata block
    metadata = [
        ("For", "Megha Poudel, Sahisa Sunuwar, Tirtha"),
        ("Agent Type", "Outbound Data Verification & Appointment Booking Agent"),
        ("Calling On Behalf Of", "Megha Poudel, Multi Dynamic Oren Park"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        set_run_font(run_label)
        run_value = p.add_run(value)
        set_run_font(run_value)

    doc.add_paragraph()  # spacer

    # --- How to Test ---
    h = doc.add_heading("How to Test", level=2)
    for run in h.runs:
        run.font.size = Pt(13)
        run.font.name = "Calibri"
        run.bold = True

    # Option 1
    p = doc.add_paragraph()
    run = p.add_run("Option 1: Receive a Test Call")
    run.bold = True
    set_run_font(run, size=11)
    p = doc.add_paragraph(
        "We will trigger a test outbound call to your mobile number. "
        "Simply answer and interact with the agent as if you were a local resident."
    )

    # Option 2
    p = doc.add_paragraph()
    run = p.add_run("Option 2: Web Demo Link")
    run.bold = True
    set_run_font(run, size=11)
    p = doc.add_paragraph(
        "Use the web demo link (provided separately) to start a conversation "
        "with the agent from your browser. No phone needed."
    )

    # --- What the Agent Does ---
    h = doc.add_heading("What the Agent Does", level=2)
    for run in h.runs:
        run.font.size = Pt(13)
        run.font.name = "Calibri"
        run.bold = True

    agent_steps = [
        (
            "Opens the call",
            "introduces itself as an AI assistant calling on behalf of Megha Poudel from Multi Dynamic Oren Park, discloses recording",
        ),
        (
            "Verifies your details",
            "confirms your name and suburb to ensure the contact record is valid",
        ),
        (
            "Asks about property plans",
            "finds out if you're thinking about buying or selling within the next twelve months",
        ),
        (
            "Categorises your interest",
            "tags you as Hot, Warm, Cold, or Unassembled based on your timeline",
        ),
        (
            "Offers a free market appraisal",
            "if you're interested in selling, offers a no-obligation property appraisal",
        ),
        (
            "Books an appointment",
            "uses the two-option method (this week or next? morning or afternoon?) to find a suitable time",
        ),
        (
            "Asks for referrals",
            "if you're not interested, asks if you know anyone who might be looking to buy or sell",
        ),
        (
            "Logs the callback",
            "if you prefer a human to call back, captures your details and notifies the team",
        ),
    ]
    for i, (bold_part, rest) in enumerate(agent_steps, 1):
        p = doc.add_paragraph(style="List Number")
        run_b = p.add_run(bold_part)
        run_b.bold = True
        set_run_font(run_b)
        run_r = p.add_run(f" \u2014 {rest}")
        set_run_font(run_r)

    # --- Test Scenarios to Try ---
    h = doc.add_heading("Test Scenarios to Try", level=2)
    for run in h.runs:
        run.font.size = Pt(13)
        run.font.name = "Calibri"
        run.bold = True

    # Table data
    table_headers = ["Scenario", "What to Say", "Expected Behaviour"]
    table_rows = [
        [
            "Happy path (seller)",
            '"Yes, I\'m thinking of selling in a few months"',
            "Agent categorises as Warm/Hot, offers free market appraisal, books appointment",
        ],
        [
            "Happy path (buyer)",
            '"I\'m looking to buy in Oren Park"',
            "Agent captures buying criteria, categorises, offers to connect with team",
        ],
        [
            "Not interested",
            '"Not interested"',
            "Agent asks one follow-up about property plans, then asks for referrals, exits gracefully",
        ],
        [
            "Already have an agent",
            '"I already have an agent"',
            "Agent acknowledges, offers second opinion, asks for referrals, ends warmly",
        ],
        [
            "Bad timing",
            '"Not a good time"',
            "Agent asks for a better time, logs callback",
        ],
        [
            "Want to speak to human",
            '"Can I speak to a real person?"',
            "Agent logs a callback request, notifies the team — does not attempt live transfer",
        ],
        [
            "Wrong number",
            '"You\'ve got the wrong person"',
            "Agent apologises, tags as invalid, ends gracefully",
        ],
        [
            "Suspicious",
            '"How did you get my number?"',
            "Agent explains local outreach, stays calm, continues if they engage",
        ],
        [
            "AI question",
            '"Are you a robot?"',
            "Agent confirms AI identity, pivots back to purpose",
        ],
        [
            "Abusive",
            "Use aggressive language",
            "Agent apologises once, then ends the call gracefully",
        ],
        [
            "Remove from list",
            '"Take me off your list"',
            "Agent confirms removal, ends immediately",
        ],
        [
            "Referral",
            '"I\'m not interested but my neighbour is selling"',
            "Agent captures the referral name and number, thanks the prospect",
        ],
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row with dark blue background and white text
    header_cells = table.rows[0].cells
    for i, header_text in enumerate(table_headers):
        cell = header_cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        # Set cell background
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="2F5496" w:val="clear"/>'
        )
        tcPr.append(shading)

    # Data rows
    for row_data in table_rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(10)
            run.font.name = "Calibri"

    # --- Providing Feedback ---
    doc.add_paragraph()  # spacer
    h = doc.add_heading("Providing Feedback", level=2)
    for run in h.runs:
        run.font.size = Pt(13)
        run.font.name = "Calibri"
        run.bold = True

    p = doc.add_paragraph()
    run = p.add_run("Use the Feedback Sheet")
    run.bold = True
    set_run_font(run)
    run = p.add_run(" (shared separately) to log:")
    set_run_font(run)

    feedback_items = [
        "What the agent said that sounded unnatural",
        "Missing information or wrong answers",
        "Scenarios that weren't handled well",
        "Suggestions for tone or wording changes",
        "Whether the lead categorisation felt accurate",
        "Whether the appointment booking flow felt smooth",
    ]
    for item in feedback_items:
        doc.add_paragraph(item, style="List Bullet")

    p = doc.add_paragraph()
    run = p.add_run("Bring your feedback to the next recurring meeting")
    run.bold = True
    set_run_font(run)
    run = p.add_run(" and we'll optimise together.")
    set_run_font(run)

    # --- Important Notes ---
    h = doc.add_heading("Important Notes", level=2)
    for run in h.runs:
        run.font.size = Pt(13)
        run.font.name = "Calibri"
        run.bold = True

    notes = [
        "This is a prototype \u2014 it will improve with your feedback and real call data",
        "The agent is set to a 10-minute max call duration for testing",
        "Operating hours are set to 8:30 AM \u2013 7:30 PM \u2014 no appointments before 11:30 AM",
        "Voicemail detection is enabled \u2014 if it detects a voicemail, it leaves a brief message",
        "All calls are recorded for review and optimisation",
        "The agent does not live transfer \u2014 it logs callbacks and sends email notifications to the team",
        "Maximum 3 call attempts per contact per day",
        "This process is being documented so it can be rolled out to other Multi Dynamic offices",
    ]
    for note in notes:
        doc.add_paragraph(note, style="List Bullet")

    # Footer
    set_footer(doc)

    # Save
    out_path = os.path.join(
        OUTPUT_DIR, "MultiDynamic_DemoInstructions_OrenPark.docx"
    )
    doc.save(out_path)
    print(f"Created: {out_path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    generate_system_prompt_docx()
    generate_demo_instructions_docx()
    print("\nDone \u2014 both files generated successfully.")
