"""
Generate two .docx files for the Direct Eight project:
1. DirectEight_SystemPrompt_Sarah.docx
2. DirectEight_DemoInstructions_Sarah.docx
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

OUTPUT_DIR = r"d:\Imp Claude Codes\Inspra Agent Skills\output"


def set_footer(doc, text="Prepared by Inspra AI"):
    """Add a centered gray footer to all sections."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.font.name = "Calibri"


def set_margins(doc, inches=1.0):
    """Set 1-inch margins on all sections."""
    for section in doc.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


def add_shaded_monospace_paragraph(doc, text):
    """Add a paragraph with Courier New font and light gray/blue shading (D9E2F3)."""
    p = doc.add_paragraph()
    # Add shading to the paragraph
    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    return p


def set_run_font(run, name="Calibri", size=11):
    run.font.name = name
    run.font.size = Pt(size)
    return run


# ---------------------------------------------------------------------------
# FILE 1: System Prompt
# ---------------------------------------------------------------------------
def generate_system_prompt_docx():
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    set_margins(doc)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("Direct Eight — Outbound Voice Agent System Prompt")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = "Calibri"

    doc.add_paragraph()  # spacer

    # Read the markdown file
    md_path = os.path.join(OUTPUT_DIR, "Direct Eight - System Prompt.md")
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    # Define the sections we want to extract
    section_names = [
        "Identity",
        "Context",
        "Conversation Principles",
        "Voice Format Rules",
        "Conversation Flow",
        "Objection Handling",
        "Tools Available",
        "Guardrails",
        "Boundaries",
    ]

    # Split content by ## headings
    # Pattern: find ## Heading and everything until next ## or end
    sections = re.split(r"^## ", content, flags=re.MULTILINE)

    for section_text in sections:
        if not section_text.strip():
            continue

        # Extract heading (first line)
        lines = section_text.strip().split("\n", 1)
        heading = lines[0].strip()

        # Skip the top-level title and horizontal rules
        if heading.startswith("#") or heading == "---":
            continue

        # Check if this heading matches one of our sections
        matched = False
        for sn in section_names:
            if heading.lower() == sn.lower():
                matched = True
                break

        if not matched:
            continue

        # Add Heading 2
        h = doc.add_heading(heading, level=2)
        for run in h.runs:
            run.font.size = Pt(13)
            run.font.name = "Calibri"
            run.bold = True

        # Add body content in monospace with shading
        body = lines[1].strip() if len(lines) > 1 else ""
        if body:
            # Split into logical paragraphs (double newline or meaningful blocks)
            paragraphs = re.split(r"\n\n+", body)
            for para_text in paragraphs:
                para_text = para_text.strip()
                if para_text:
                    add_shaded_monospace_paragraph(doc, para_text)

    # Footer
    set_footer(doc)

    # Save
    out_path = os.path.join(OUTPUT_DIR, "DirectEight_SystemPrompt_Sarah.docx")
    doc.save(out_path)
    print(f"Created: {out_path}")


# ---------------------------------------------------------------------------
# FILE 2: Demo Instructions
# ---------------------------------------------------------------------------
def generate_demo_instructions_docx():
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    set_margins(doc)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("Direct Eight — Prototype Demo Instructions")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = "Calibri"

    doc.add_paragraph()  # spacer

    # Metadata block
    metadata = [
        ("For", "Simon & Prakash"),
        ("Agent Name", "Sarah (AI Business Development Rep)"),
        ("Type", "Outbound Cold Calling Agent"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        set_run_font(run_label)
        run_value = p.add_run(value)
        set_run_font(run_value)

    doc.add_paragraph()  # spacer

    # Read the markdown file
    md_path = os.path.join(OUTPUT_DIR, "Direct Eight - Demo Instructions.md")
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    # --- How to Test ---
    h = doc.add_heading("How to Test", level=2)
    for run in h.runs:
        run.font.size = Pt(13)
        run.font.name = "Calibri"
        run.bold = True

    # Option 1
    p = doc.add_paragraph()
    run = p.add_run("Option 1: Receive a Test Call")
    run.bold = True
    set_run_font(run, size=11)
    p = doc.add_paragraph(
        "We will trigger a test outbound call to your mobile number. "
        "Simply answer and interact with Sarah as if you were a prospect."
    )

    # Option 2
    p = doc.add_paragraph()
    run = p.add_run("Option 2: Web Demo Link")
    run.bold = True
    set_run_font(run, size=11)
    p = doc.add_paragraph(
        "Use the web demo link (provided separately) to start a conversation "
        "with Sarah from your browser. No phone needed."
    )

    # --- What Sarah Does ---
    h = doc.add_heading("What Sarah Does", level=2)
    for run in h.runs:
        run.font.size = Pt(13)
        run.font.name = "Calibri"
        run.bold = True

    sarah_steps = [
        ("Opens the call", "introduces herself, discloses she's an AI assistant, and mentions Direct Eight"),
        ("Asks discovery questions", "finds out who handles cleaning, what's working, what's not"),
        ("Presents the offer", "connects your pain points to Direct Eight's services (CleanPlan, in-house team, sanitation expertise)"),
        ("Books a meeting", "offers a 15-minute consultation via Google Calendar with a Zoom link"),
        ("Live transfers", "if you're very keen, she can connect you to the team directly"),
        ("Logs callbacks", "if you prefer a human to call back, she captures your details and notifies the team"),
    ]
    for i, (bold_part, rest) in enumerate(sarah_steps, 1):
        p = doc.add_paragraph(style="List Number")
        run_b = p.add_run(bold_part)
        run_b.bold = True
        set_run_font(run_b)
        run_r = p.add_run(f" — {rest}")
        set_run_font(run_r)

    # --- Test Scenarios to Try ---
    h = doc.add_heading("Test Scenarios to Try", level=2)
    for run in h.runs:
        run.font.size = Pt(13)
        run.font.name = "Calibri"
        run.bold = True

    # Table data
    table_headers = ["Scenario", "What to Say", "Expected Behaviour"]
    table_rows = [
        ["Happy path", '"Yes, tell me more" → agree to a meeting', "Sarah qualifies, pitches, books a meeting"],
        ["Not interested", '"Not interested"', "Sarah asks one follow-up, then exits gracefully"],
        ["Already have a cleaner", '"We already have someone"', "Sarah asks about consistency, doesn't push"],
        ["Bad timing", '"Not a good time"', "Sarah asks for a better time, logs callback"],
        ["Want to speak to human", '"Can I speak to a real person?"', "Sarah transfers or logs a callback"],
        ["Price concern", '"How much does it cost?"', "Sarah redirects to free consultation, never quotes pricing"],
        ["Suspicious", '"Is this a scam?"', "Sarah stays calm, mentions Collins Street office, offers website"],
        ["AI question", '"Are you a robot?"', "Sarah confirms AI, pivots back to purpose"],
        ["Abusive", "Use aggressive language", "Sarah apologises once, then ends the call"],
        ["Remove from list", '"Take me off your list"', "Sarah confirms removal, ends immediately"],
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row with dark blue background and white text
    header_cells = table.rows[0].cells
    for i, header_text in enumerate(table_headers):
        cell = header_cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        # Set cell background
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2F5496" w:val="clear"/>')
        tcPr.append(shading)

    # Data rows
    for row_data in table_rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(10)
            run.font.name = "Calibri"

    # --- Providing Feedback ---
    doc.add_paragraph()  # spacer
    h = doc.add_heading("Providing Feedback", level=2)
    for run in h.runs:
        run.font.size = Pt(13)
        run.font.name = "Calibri"
        run.bold = True

    p = doc.add_paragraph()
    run = p.add_run("Use the Feedback Sheet")
    run.bold = True
    set_run_font(run)
    run = p.add_run(" (shared separately) to log:")
    set_run_font(run)

    feedback_items = [
        "What the bot said that sounded unnatural",
        "Missing information or wrong answers",
        "Scenarios that weren't handled well",
        "Suggestions for tone or wording changes",
    ]
    for item in feedback_items:
        p = doc.add_paragraph(item, style="List Bullet")

    p = doc.add_paragraph()
    run = p.add_run("Bring your feedback to our next Tuesday meeting")
    run.bold = True
    set_run_font(run)
    run = p.add_run(" and we'll optimise together.")
    set_run_font(run)

    # --- Important Notes ---
    h = doc.add_heading("Important Notes", level=2)
    for run in h.runs:
        run.font.size = Pt(13)
        run.font.name = "Calibri"
        run.bold = True

    notes = [
        "This is a prototype — it will improve with your feedback and real call data",
        "The bot is set to a 5-minute max call duration for testing",
        "Voicemail detection is enabled — if it detects a voicemail, it leaves a brief message",
        "All calls are recorded for review and optimisation",
        "The live transfer number is currently set to the emergency response line — we'll update this to the dedicated sales line once confirmed",
    ]
    for note in notes:
        p = doc.add_paragraph(note, style="List Bullet")

    # Footer
    set_footer(doc)

    # Save
    out_path = os.path.join(OUTPUT_DIR, "DirectEight_DemoInstructions_Sarah.docx")
    doc.save(out_path)
    print(f"Created: {out_path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    generate_system_prompt_docx()
    generate_demo_instructions_docx()
    print("\nDone — both files generated successfully.")
