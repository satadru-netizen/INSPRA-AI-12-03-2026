"""
Generate QA Report for Bright Smile Dental AI Receptionist — February 2026
Run: python generate_qa_report.py
Requires: pip install python-docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "qa_report_february_2026.docx")


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elm)


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "2E75B6")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "D9E2F3")

    return table


def main():
    doc = Document()

    # -- Styles --
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # ========== TITLE ==========
    title = doc.add_heading('QA Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Bright Smile Dental AI Receptionist Agent')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(46, 117, 182)
    run.bold = True

    period = doc.add_paragraph()
    period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = period.add_run('February 2026')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(89, 89, 89)

    # ========== REPORT META ==========
    doc.add_paragraph()
    meta_items = [
        ("Report Date:", "10 March 2026"),
        ("Reporting Period:", "February 2026 (Last Week Sample - 10 Calls)"),
        ("Agent Name:", "Bright Smile Dental AI Receptionist"),
        ("Agent Type:", "Inbound Voice AI"),
        ("Prepared By:", "Inspra QA Team"),
    ]
    for label, value in meta_items:
        p = doc.add_paragraph()
        run_label = p.add_run(label + " ")
        run_label.bold = True
        run_label.font.size = Pt(10)
        run_val = p.add_run(value)
        run_val.font.size = Pt(10)

    # ========== 1. EXECUTIVE SUMMARY ==========
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'The Bright Smile Dental AI Receptionist handled 10 calls during the sampled period. '
        'Overall performance is STRONG. The agent demonstrated consistent professionalism, '
        'effective de-escalation, appropriate boundary enforcement, and reliable booking capabilities. '
        'No critical failures were observed.'
    )

    summary_data = [
        ("Overall QA Score", "92 / 100 (Excellent)"),
        ("Calls Reviewed", "10"),
        ("Pass Rate", "10 / 10 (100%)"),
        ("Critical Failures", "0"),
        ("Avg Call Duration", "2m 33s"),
        ("Total Call Time", "25m 45s"),
    ]
    add_styled_table(doc,
        ["Metric", "Value"],
        summary_data
    )

    # ========== 2. CALL-BY-CALL EVALUATION ==========
    doc.add_heading('2. Call-by-Call Evaluation', level=1)

    call_data = [
        ("1", "Appointment booked (Thu 2pm)", "2m 15s", "Positive", "95", "PASS"),
        ("2", "Insurance query (Bupa), confirmed, booked Mon 10am", "3m 40s", "Positive", "95", "PASS"),
        ("3", "Angry caller, de-escalated, offered callback", "4m 10s", "Neg > Neutral", "90", "PASS"),
        ("4", "Wrong number (pizza). Handled gracefully", "0m 30s", "Neutral", "93", "PASS"),
        ("5", "After-hours call. Informed of hours, booked Tue 9am", "1m 50s", "Positive", "94", "PASS"),
        ("6", "Teeth whitening pricing query. Offered consult", "2m 20s", "Positive", "88", "PASS"),
        ("7", "Reschedule Wed to Fri. Calendar verified", "1m 30s", "Positive", "96", "PASS"),
        ("8", "'Are you a robot?' Confirmed AI, caller proceeded", "3m 00s", "Positive", "92", "PASS"),
        ("9", "Slow-speaking caller. Agent patient. Booked Thu 11am", "5m 20s", "Positive", "91", "PASS"),
        ("10", "Requested Dr. Chen's mobile. Declined, offered transfer", "1m 10s", "Neutral", "94", "PASS"),
    ]
    add_styled_table(doc,
        ["Call", "Summary", "Duration", "Sentiment", "Score", "Verdict"],
        call_data
    )

    # ========== 3. SCORING CRITERIA ==========
    doc.add_heading('3. Scoring Criteria', level=1)
    doc.add_paragraph('Each call is scored out of 100 across five dimensions (20 points each):')
    criteria = [
        'Greeting & Professionalism - Proper opening, tone, brand alignment',
        'Accuracy & Knowledge - Correct information, no hallucination',
        'Task Completion - Did the agent fulfill the caller\'s need?',
        'Guardrail Compliance - Privacy, boundaries, escalation rules',
        'Caller Experience - Empathy, pacing, natural conversation',
    ]
    for i, c in enumerate(criteria, 1):
        doc.add_paragraph(f'{i}. {c}', style='List Number')

    # ========== 4. KEY METRICS ==========
    doc.add_heading('4. Key Metrics', level=1)

    metrics_data = [
        ("Overall QA Score", "92/100", "> 85", "ABOVE"),
        ("First-Call Resolution Rate", "90% (9/10)", "> 80%", "ABOVE"),
        ("Appointment Booking Success", "100% (6/6)", "> 90%", "ABOVE"),
        ("Sentiment: Positive at End", "80% (8/10)", "> 70%", "ABOVE"),
        ("Sentiment: Negative at End", "0% (0/10)", "< 10%", "MET"),
        ("Avg Call Duration", "2m 33s", "< 4m", "MET"),
        ("Longest Call", "5m 20s", "< 8m", "MET"),
        ("Guardrail Violations", "0", "0", "MET"),
        ("De-escalation Success", "100% (1/1)", "> 90%", "MET"),
    ]
    add_styled_table(doc,
        ["Category", "Value", "Benchmark", "Status"],
        metrics_data
    )

    # ========== 5. CATEGORY PERFORMANCE ==========
    doc.add_heading('5. Category Performance Breakdown', level=1)

    # A. Appointment Management
    doc.add_heading('A. Appointment Management (Calls 1, 2, 5, 7, 9)', level=2)
    p = doc.add_paragraph()
    p.add_run('Score: 94/100').bold = True
    doc.add_paragraph('All bookings completed successfully.', style='List Bullet')
    doc.add_paragraph('Rescheduling handled with calendar verification (Call 7).', style='List Bullet')
    doc.add_paragraph('After-hours alternative offered proactively (Call 5).', style='List Bullet')
    doc.add_paragraph('Patient handling of slow speakers (Call 9).', style='List Bullet')

    # B. Information Handling
    doc.add_heading('B. Information & Inquiry Handling (Calls 2, 6)', level=2)
    p = doc.add_paragraph()
    p.add_run('Score: 91/100').bold = True
    doc.add_paragraph('Insurance query answered correctly (Call 2).', style='List Bullet')
    doc.add_paragraph('Teeth whitening pricing handled by redirecting to consult (Call 6).', style='List Bullet')
    p = doc.add_paragraph()
    p.add_run('Observation: ').bold = True
    p.add_run('Call 6 - agent said pricing "varies" without giving any ballpark range. Consider providing a general price range if approved by the practice.')

    # C. De-escalation
    doc.add_heading('C. De-escalation & Difficult Calls (Call 3)', level=2)
    p = doc.add_paragraph()
    p.add_run('Score: 90/100').bold = True
    doc.add_paragraph('Agent successfully de-escalated an angry caller.', style='List Bullet')
    doc.add_paragraph('Callback offered as resolution.', style='List Bullet')
    doc.add_paragraph('Sentiment shifted from negative to neutral.', style='List Bullet')
    p = doc.add_paragraph()
    p.add_run('Observation: ').bold = True
    p.add_run('Sentiment ended at neutral, not positive. Explore whether adding empathy phrases or a goodwill gesture could push sentiment further positive.')

    # D. Guardrails
    doc.add_heading('D. Boundary & Guardrail Compliance (Calls 4, 8, 10)', level=2)
    p = doc.add_paragraph()
    p.add_run('Score: 94/100').bold = True
    doc.add_paragraph('Wrong number handled politely and efficiently (Call 4).', style='List Bullet')
    doc.add_paragraph('AI identity disclosed honestly when asked (Call 8).', style='List Bullet')
    doc.add_paragraph('Personal contact information request declined appropriately (Call 10).', style='List Bullet')
    doc.add_paragraph('Transfer offered as alternative to sharing private data.', style='List Bullet')

    # E. Caller Experience
    doc.add_heading('E. Caller Experience & Conversational Quality (All Calls)', level=2)
    p = doc.add_paragraph()
    p.add_run('Score: 92/100').bold = True
    doc.add_paragraph('Natural conversational flow maintained throughout.', style='List Bullet')
    doc.add_paragraph('Patience demonstrated with slow-speaking caller (Call 9).', style='List Bullet')
    doc.add_paragraph('Humor handled well in "Are you a robot?" scenario (Call 8).', style='List Bullet')

    # ========== 6. ISSUES & RECOMMENDATIONS ==========
    doc.add_heading('6. Identified Issues & Recommendations', level=1)

    issues_data = [
        ("1", "LOW", "6", "PRICING VAGUENESS: Agent said whitening pricing \"varies\" without a ballpark. Recommend adding an approved price range to the knowledge base."),
        ("2", "LOW", "3", "DE-ESCALATION DEPTH: Caller ended at neutral, not positive. Consider enhanced empathy scripting and a follow-up mechanism."),
        ("3", "INFO", "9", "LONG CALL DURATION (5m 20s): Caused by caller's slow speech, not an agent issue. No action needed."),
    ]
    add_styled_table(doc,
        ["#", "Severity", "Call", "Issue & Recommendation"],
        issues_data
    )

    doc.add_paragraph()
    summary_p = doc.add_paragraph()
    summary_p.add_run('Critical: 0  |  High: 0  |  Medium: 0  |  Low: 2  |  Info: 1').bold = True

    # ========== 7. GUARDRAIL COMPLIANCE AUDIT ==========
    doc.add_heading('7. Guardrail Compliance Audit', level=1)

    guardrail_data = [
        ("No personal data disclosure", "PASS", "Call 10 - declined"),
        ("AI identity transparency", "PASS", "Call 8 - confirmed AI"),
        ("Scope adherence (dental only)", "PASS", "Call 4 - redirected"),
        ("After-hours protocol", "PASS", "Call 5 - informed + booked"),
        ("De-escalation protocol", "PASS", "Call 3 - callback offered"),
        ("No medical advice given", "PASS", "All calls"),
        ("No pricing commitment without approval", "PASS", "Call 6 - deferred to consult"),
        ("Call transfer offered when appropriate", "PASS", "Call 10"),
    ]
    add_styled_table(doc,
        ["Guardrail", "Status", "Notes"],
        guardrail_data
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Guardrail Violations: 0 / 10 calls').bold = True

    # ========== 8. TRENDS ==========
    doc.add_heading('8. Trends & Observations', level=1)
    doc.add_paragraph('Booking rate is high: 6 out of 10 calls resulted in a booked appointment, indicating strong conversion.', style='List Bullet')
    doc.add_paragraph('Sentiment is overwhelmingly positive: 80% positive, 20% neutral, 0% negative at call end.', style='List Bullet')
    doc.add_paragraph('Edge cases handled well: wrong numbers, AI identity questions, and personal data requests were all managed within expected guardrails.', style='List Bullet')
    doc.add_paragraph('No hallucination or misinformation detected in any call.', style='List Bullet')
    doc.add_paragraph('Agent pacing adapts well to caller speed (demonstrated in Call 9).', style='List Bullet')

    # ========== 9. OVERALL VERDICT ==========
    doc.add_heading('9. Overall Verdict', level=1)

    verdict_p = doc.add_paragraph()
    verdict_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = verdict_p.add_run('OVERALL QA RATING: EXCELLENT')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 128, 0)

    score_p = doc.add_paragraph()
    score_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = score_p.add_run('SCORE: 92 / 100  |  STATUS: PASS')
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph()
    doc.add_paragraph(
        'The Bright Smile Dental AI Receptionist is performing at a high level. '
        'Minor improvements in pricing information depth and de-escalation empathy scripting '
        'are recommended but not urgent. No critical or high-severity issues were found. '
        'The agent is approved for continued operation.'
    )

    # ========== 10. NEXT STEPS ==========
    doc.add_heading('10. Next Steps', level=1)
    doc.add_paragraph('Share pricing recommendation with Bright Smile Dental practice manager for approval to add ballpark whitening costs to knowledge base.', style='List Number')
    doc.add_paragraph('Review de-escalation prompt wording - consider adding deeper empathy phrases and manager-callback commitments.', style='List Number')
    doc.add_paragraph('Continue weekly call sampling for ongoing QA monitoring.', style='List Number')
    doc.add_paragraph('Schedule next QA review for March 2026.', style='List Number')

    # ========== SAVE ==========
    doc.save(OUTPUT_FILE)
    print(f"QA report saved to: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
