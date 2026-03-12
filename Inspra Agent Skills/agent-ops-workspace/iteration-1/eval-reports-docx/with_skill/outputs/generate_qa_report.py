"""Generate QA Report for Bright Smile Dental - February 2026"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Call Data & Scoring ──────────────────────────────────────────────────────

calls = [
    {
        "id": 1, "summary": "Caller booked appointment for Thursday 2pm.",
        "duration": "2m 15s", "sentiment": "Positive",
        "greeting": 5, "listening": 4, "accuracy": 5, "tone": 5, "flow": 5, "compliance": 5, "resolution": 5,
        "notes": {
            "greeting": "Professional opening, identified clinic.",
            "listening": "Good listening; standard booking call with minimal back-and-forth.",
            "accuracy": "Correct date/time confirmed.",
            "tone": "Warm and friendly throughout.",
            "flow": "Seamless booking flow, no unnecessary steps.",
            "compliance": "No boundary issues.",
            "resolution": "Appointment booked successfully, next steps confirmed.",
        }
    },
    {
        "id": 2, "summary": "Insurance query (Bupa) + booking for Monday 10am.",
        "duration": "3m 40s", "sentiment": "Positive",
        "greeting": 5, "listening": 5, "accuracy": 5, "tone": 5, "flow": 5, "compliance": 5, "resolution": 5,
        "notes": {
            "greeting": "Warm greeting, identified self and clinic.",
            "listening": "Answered insurance question first before moving to booking.",
            "accuracy": "Correctly confirmed Bupa acceptance.",
            "tone": "Friendly and reassuring on insurance topic.",
            "flow": "Natural transition from insurance query to booking.",
            "compliance": "No unauthorized commitments about coverage details.",
            "resolution": "Both insurance question and booking handled perfectly.",
        }
    },
    {
        "id": 3, "summary": "Angry caller, agent de-escalated and offered callback.",
        "duration": "4m 10s", "sentiment": "Started negative, ended neutral",
        "greeting": 4, "listening": 5, "accuracy": 4, "tone": 5, "flow": 4, "compliance": 5, "resolution": 4,
        "notes": {
            "greeting": "Professional but caller was already agitated; slightly limited opening.",
            "listening": "Excellent — let the caller vent without interrupting, acknowledged frustration.",
            "accuracy": "No incorrect info given; appropriately vague on wait-time causes.",
            "tone": "Outstanding de-escalation; empathetic and patient throughout.",
            "flow": "Minor hesitation before offering callback, but recovered well.",
            "compliance": "Stayed within boundaries, no unauthorized promises.",
            "resolution": "Did not fully resolve original issue but offered appropriate callback alternative.",
        }
    },
    {
        "id": 4, "summary": "Wrong number — caller wanted pizza delivery.",
        "duration": "30s", "sentiment": "Neutral",
        "greeting": 5, "listening": 4, "accuracy": 5, "tone": 5, "flow": 5, "compliance": 5, "resolution": 5,
        "notes": {
            "greeting": "Proper greeting that immediately clarified the business.",
            "listening": "Quickly identified the mismatch.",
            "accuracy": "Correctly identified this is a dental clinic, not pizza.",
            "tone": "Handled gracefully with humor/warmth.",
            "flow": "Efficient — wrapped up quickly without dragging on.",
            "compliance": "No issues.",
            "resolution": "Correct resolution — informed caller of wrong number politely.",
        }
    },
    {
        "id": 5, "summary": "After-hours call; informed of hours, booked for Tuesday 9am.",
        "duration": "1m 50s", "sentiment": "Positive",
        "greeting": 5, "listening": 4, "accuracy": 5, "tone": 4, "flow": 5, "compliance": 5, "resolution": 5,
        "notes": {
            "greeting": "Clear greeting including after-hours notice.",
            "listening": "Good, addressed caller's need promptly.",
            "accuracy": "Correct business hours communicated, accurate booking.",
            "tone": "Professional; slightly less warm given after-hours automation feel.",
            "flow": "Smooth — hours info then booking offer in logical sequence.",
            "compliance": "Proper disclosure of after-hours status.",
            "resolution": "Booking completed for next available slot.",
        }
    },
    {
        "id": 6, "summary": "Teeth whitening pricing inquiry; offered consult booking.",
        "duration": "2m 20s", "sentiment": "Positive",
        "greeting": 5, "listening": 4, "accuracy": 4, "tone": 5, "flow": 4, "compliance": 5, "resolution": 4,
        "notes": {
            "greeting": "Warm, professional opening.",
            "listening": "Addressed pricing question before pivoting to booking.",
            "accuracy": "Said pricing 'varies' — correct but could have given a range if authorized.",
            "tone": "Enthusiastic and encouraging about the consultation.",
            "flow": "Slight awkwardness in pivot from 'I can't give exact price' to 'book a consult'.",
            "compliance": "Correctly avoided quoting unauthorized prices.",
            "resolution": "Offered consult but unclear if caller actually booked.",
        }
    },
    {
        "id": 7, "summary": "Rescheduled from Wednesday to Friday.",
        "duration": "1m 30s", "sentiment": "Positive",
        "greeting": 5, "listening": 5, "accuracy": 5, "tone": 5, "flow": 5, "compliance": 5, "resolution": 5,
        "notes": {
            "greeting": "Friendly opening.",
            "listening": "Immediately understood the reschedule request.",
            "accuracy": "Checked calendar correctly, confirmed new slot.",
            "tone": "Accommodating and pleasant.",
            "flow": "Quick and efficient reschedule process.",
            "compliance": "No issues.",
            "resolution": "Reschedule confirmed with clear next steps.",
        }
    },
    {
        "id": 8, "summary": "Caller asked 'Are you a robot?' — agent confirmed AI, caller proceeded to book.",
        "duration": "3m", "sentiment": "Positive",
        "greeting": 5, "listening": 5, "accuracy": 5, "tone": 5, "flow": 5, "compliance": 5, "resolution": 5,
        "notes": {
            "greeting": "Standard warm greeting.",
            "listening": "Handled the 'robot' question naturally without deflecting.",
            "accuracy": "Honestly confirmed AI status — no deception.",
            "tone": "Light and natural; caller laughed, good rapport.",
            "flow": "Smooth transition from AI disclosure to booking.",
            "compliance": "Perfect — disclosed AI status when asked (required).",
            "resolution": "Booking completed successfully.",
        }
    },
    {
        "id": 9, "summary": "Slow-speaking caller with pauses; agent stayed patient. Booked Thursday 11am.",
        "duration": "5m 20s", "sentiment": "Positive",
        "greeting": 5, "listening": 5, "accuracy": 5, "tone": 5, "flow": 4, "compliance": 5, "resolution": 5,
        "notes": {
            "greeting": "Warm greeting.",
            "listening": "Excellent patience — never rushed the caller, waited through long pauses.",
            "accuracy": "All booking details confirmed correctly.",
            "tone": "Patient and accommodating throughout the longer call.",
            "flow": "Slightly longer pauses made flow feel slower, but agent adapted well.",
            "compliance": "No issues.",
            "resolution": "Appointment booked with confirmation.",
        }
    },
    {
        "id": 10, "summary": "Caller requested Dr. Chen's personal mobile. Agent declined, offered transfer.",
        "duration": "1m 10s", "sentiment": "Neutral",
        "greeting": 5, "listening": 4, "accuracy": 5, "tone": 4, "flow": 5, "compliance": 5, "resolution": 4,
        "notes": {
            "greeting": "Professional opening.",
            "listening": "Understood request quickly.",
            "accuracy": "Correctly identified this as a request to decline.",
            "tone": "Professional but could have been slightly warmer in the refusal.",
            "flow": "Efficient handling — decline then immediate alternative.",
            "compliance": "Excellent — correctly refused to share personal contact info.",
            "resolution": "Offered transfer as alternative; unclear if caller accepted.",
        }
    },
]

# ── Weighted Score Calculation ───────────────────────────────────────────────

weights = {
    "greeting": 0.10, "listening": 0.15, "accuracy": 0.20,
    "tone": 0.15, "flow": 0.15, "compliance": 0.10, "resolution": 0.15
}

dimension_labels = {
    "greeting": "Greeting & ID",
    "listening": "Listening",
    "accuracy": "Accuracy",
    "tone": "Tone & Empathy",
    "flow": "Flow & Efficiency",
    "compliance": "Compliance",
    "resolution": "Resolution",
}

for call in calls:
    raw = sum(call[d] * w for d, w in weights.items())
    call["weighted_score"] = round((raw / 5) * 100, 1)

overall_score = round(sum(c["weighted_score"] for c in calls) / len(calls), 1)

# Per-dimension averages
dim_averages = {}
for dim in weights:
    dim_averages[dim] = round(sum(c[dim] for c in calls) / len(calls), 2)

# Flagged calls: any dimension <= 2 or total < 60
flagged_calls = [c for c in calls if c["weighted_score"] < 60 or any(c[d] <= 2 for d in weights)]

# ── Document Creation ────────────────────────────────────────────────────────

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Set margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def set_cell_shading(cell, color_hex):
    """Apply shading to a table cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), color_hex)
    tc_pr.append(shading_elm)


def style_header_row(table):
    """Apply dark blue shading and white text to header row."""
    for cell in table.rows[0].cells:
        set_cell_shading(cell, '2F5496')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(10)


def add_table_row(table, values):
    """Add a row to a table with the given values."""
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = str(val)
        for paragraph in cell.paragraphs:
            paragraph.style.font.size = Pt(10)
            paragraph.style.font.name = 'Calibri'


def add_paragraph_shaded(doc, text, shade_color='D9E2F3'):
    """Add a paragraph with light shading (for transcript excerpts)."""
    para = doc.add_paragraph(text)
    para.style.font.size = Pt(10)
    para.style.font.name = 'Calibri'
    # Add shading
    pPr = para._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), shade_color)
    pPr.append(shd)
    # Add border
    pBdr = OxmlElement('w:pBdr')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '1')
        border.set(qn('w:color'), '8FAADC')
        pBdr.append(border)
    pPr.append(pBdr)
    return para


# ── Title ────────────────────────────────────────────────────────────────────

title = doc.add_heading('Bright Smile Dental — QA Audit Report', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(47, 84, 150)

# Report metadata
meta_info = [
    ("Prepared by", "Inspra AI"),
    ("Period", "1 February 2026 — 28 February 2026"),
    ("Agent", "Bright Smile Dental AI Receptionist"),
    ("Calls Audited", "10 of 10"),
]
for label, value in meta_info:
    para = doc.add_paragraph()
    run_label = para.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(11)
    run_label.font.name = 'Calibri'
    run_value = para.add_run(value)
    run_value.font.size = Pt(11)
    run_value.font.name = 'Calibri'
    para.paragraph_format.space_after = Pt(2)

doc.add_paragraph()  # spacer

# ── Summary Score (prominent) ────────────────────────────────────────────────

score_heading = doc.add_heading('Summary Score', level=2)
for run in score_heading.runs:
    run.font.size = Pt(13)
    run.font.bold = True

score_para = doc.add_paragraph()
score_run = score_para.add_run(f"{overall_score} / 100")
score_run.font.size = Pt(20)
score_run.font.bold = True
score_run.font.color.rgb = RGBColor(47, 84, 150)
score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
# Shade the score paragraph
pPr = score_para._element.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear')
shd.set(qn('w:color'), 'auto')
shd.set(qn('w:fill'), 'D9E2F3')
pPr.append(shd)

rating_para = doc.add_paragraph()
if overall_score >= 90:
    rating_text = "Excellent — The agent is performing at a high standard across all quality dimensions."
elif overall_score >= 80:
    rating_text = "Good — The agent performs well with minor areas for improvement."
elif overall_score >= 70:
    rating_text = "Satisfactory — Functional but several areas need attention."
else:
    rating_text = "Needs Improvement — Significant quality gaps identified."
rating_run = rating_para.add_run(rating_text)
rating_run.font.size = Pt(11)
rating_run.font.name = 'Calibri'
rating_run.italic = True
rating_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Score Breakdown Table ────────────────────────────────────────────────────

doc.add_heading('Score Breakdown', level=2)
for run in doc.paragraphs[-1].runs:
    run.font.size = Pt(13)
    run.font.bold = True

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Dimension', 'Avg Score (out of 5)', 'Weight', 'Weighted Contribution']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

style_header_row(table)

for dim, weight in weights.items():
    avg = dim_averages[dim]
    contribution = round(avg * weight * 20, 1)  # (avg/5)*weight*100
    add_table_row(table, [
        dimension_labels[dim],
        f"{avg}",
        f"{int(weight * 100)}%",
        f"{contribution}",
    ])

# Total row
total_row = table.add_row()
total_row.cells[0].text = "TOTAL"
total_row.cells[1].text = ""
total_row.cells[2].text = "100%"
total_row.cells[3].text = f"{overall_score}"
for cell in total_row.cells:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True

# ── Individual Call Scores ───────────────────────────────────────────────────

doc.add_heading('Individual Call Scores', level=2)
for run in doc.paragraphs[-1].runs:
    run.font.size = Pt(13)
    run.font.bold = True

call_table = doc.add_table(rows=1, cols=10)
call_table.style = 'Table Grid'
call_table.alignment = WD_TABLE_ALIGNMENT.CENTER

call_headers = ['Call', 'Greet', 'Listen', 'Accur', 'Tone', 'Flow', 'Compl', 'Resol', 'Score', 'Sentiment']
for i, header in enumerate(call_headers):
    cell = call_table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

style_header_row(call_table)

for c in calls:
    row = call_table.add_row()
    values = [
        f"#{c['id']}", str(c['greeting']), str(c['listening']), str(c['accuracy']),
        str(c['tone']), str(c['flow']), str(c['compliance']), str(c['resolution']),
        f"{c['weighted_score']}", c['sentiment']
    ]
    for i, val in enumerate(values):
        row.cells[i].text = val
        row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Flagged Calls ────────────────────────────────────────────────────────────

doc.add_heading('Flagged Calls', level=2)
for run in doc.paragraphs[-1].runs:
    run.font.size = Pt(13)
    run.font.bold = True

if not flagged_calls:
    para = doc.add_paragraph()
    run = para.add_run("No calls were flagged for review this period. All calls scored above 60/100 and no individual dimension scored 2 or below.")
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.italic = True
else:
    for c in flagged_calls:
        doc.add_heading(f"Call #{c['id']} — Score: {c['weighted_score']}/100", level=3)
        para = doc.add_paragraph()
        issue_label = para.add_run("Issue: ")
        issue_label.bold = True
        para.add_run(c['summary'])

# ── Calls Requiring Closer Attention (not flagged but worth noting) ──────────

doc.add_heading('Calls Requiring Closer Attention', level=2)
for run in doc.paragraphs[-1].runs:
    run.font.size = Pt(13)
    run.font.bold = True

# Call 3 — de-escalation call
doc.add_heading('Call #3 — De-escalation (Angry Caller)', level=3)
para = doc.add_paragraph()
para.add_run("Score: ").bold = True
para.add_run(f"{calls[2]['weighted_score']}/100")
para = doc.add_paragraph()
para.add_run("Context: ").bold = True
para.add_run("Caller was angry about previous hold time. Agent de-escalated successfully and offered a callback.")
para = doc.add_paragraph()
para.add_run("Observation: ").bold = True
para.add_run("The agent handled this well overall. The greeting was slightly constrained by the caller's agitation, and there was a minor hesitation before offering the callback option. Sentiment shifted from negative to neutral — a positive outcome for a difficult call.")
add_paragraph_shaded(doc,
    "Caller: \"I've been waiting for 20 minutes last time, this is ridiculous!\"\n"
    "Agent: \"I completely understand your frustration, and I'm sorry about that experience. "
    "Let me make sure we take care of you right away...\"")
para = doc.add_paragraph()
para.add_run("Recommended Improvement: ").bold = True
para.add_run("Consider adding a prompt instruction for the agent to proactively offer a callback or priority rebooking earlier in de-escalation scenarios, rather than waiting for a natural pause.")

# Call 6 — pricing inquiry
doc.add_heading('Call #6 — Pricing Inquiry (Teeth Whitening)', level=3)
para = doc.add_paragraph()
para.add_run("Score: ").bold = True
para.add_run(f"{calls[5]['weighted_score']}/100")
para = doc.add_paragraph()
para.add_run("Context: ").bold = True
para.add_run("Caller asked about teeth whitening pricing. Agent said it varies and offered to book a consultation.")
para = doc.add_paragraph()
para.add_run("Observation: ").bold = True
para.add_run("The agent correctly avoided quoting unauthorized prices (strong compliance), but the transition from 'I can't give a price' to 'book a consult' felt slightly abrupt. Accuracy was docked slightly because a general price range could have been provided if authorized. It's also unclear whether the caller actually booked the consultation.")
add_paragraph_shaded(doc,
    "Caller: \"How much does teeth whitening cost?\"\n"
    "Agent: \"Great question! The cost can vary depending on the treatment option "
    "that's best for you. I'd recommend booking a consultation so the dentist can "
    "give you an accurate quote. Would you like to schedule one?\"")
para = doc.add_paragraph()
para.add_run("Recommended Improvement: ").bold = True
para.add_run("If authorized, equip the agent with general price ranges (e.g., 'Teeth whitening typically starts from $X') to give callers a ballpark before offering the consult. Also add a confirmation step to track whether the consult was actually booked.")

# Call 10 — personal info request
doc.add_heading('Call #10 — Personal Information Request', level=3)
para = doc.add_paragraph()
para.add_run("Score: ").bold = True
para.add_run(f"{calls[9]['weighted_score']}/100")
para = doc.add_paragraph()
para.add_run("Context: ").bold = True
para.add_run("Caller asked for Dr. Chen's personal mobile number. Agent declined and offered to transfer.")
para = doc.add_paragraph()
para.add_run("Observation: ").bold = True
para.add_run("Excellent compliance — the agent correctly refused to share personal contact information. The tone could have been slightly warmer during the refusal to maintain rapport. The resolution is marked down slightly because it's unclear if the caller accepted the transfer offer.")
add_paragraph_shaded(doc,
    "Caller: \"Can I get Dr. Chen's mobile number?\"\n"
    "Agent: \"I'm not able to share personal contact details, but I can transfer "
    "you to the clinic line where the team can help you reach Dr. Chen. "
    "Would you like me to do that?\"")
para = doc.add_paragraph()
para.add_run("Recommended Improvement: ").bold = True
para.add_run("Add softer framing to the refusal, e.g., 'For privacy reasons, I'm not able to share personal numbers, but I'd love to help you connect with Dr. Chen through the clinic.' Also ensure the transfer is tracked to confirm resolution.")

# ── Patterns & Trends ───────────────────────────────────────────────────────

doc.add_heading('Patterns & Trends', level=2)
for run in doc.paragraphs[-1].runs:
    run.font.size = Pt(13)
    run.font.bold = True

trends = [
    ("Strong greeting consistency", "All 10 calls scored 4 or 5 on Greeting & ID. The agent consistently identifies itself and the clinic, creating a professional first impression."),
    ("Excellent compliance", "Perfect 5/5 compliance scores across all calls. The agent never shared unauthorized information, made improper commitments, or failed to disclose its AI status when asked."),
    ("Outstanding patience with diverse callers", "From the slow-speaking caller (#9) to the angry caller (#3) to the wrong-number caller (#4), the agent adapted its pace and approach appropriately."),
    ("Pricing inquiries are a weak point", "Call #6 revealed a gap in how pricing questions are handled. The agent correctly avoids quoting prices but the pivot to 'book a consult' needs smoother framing."),
    ("High booking conversion", "7 out of 10 calls resulted in a booking or reschedule — a 70% conversion rate on all calls (87.5% when excluding the wrong number and personal info request calls)."),
    ("Resolution tracking gap", "In calls #6 and #10, it's unclear whether the offered next step (consult booking, transfer) was actually completed. This is a tracking/analytics gap rather than an agent quality issue."),
]

for title_text, detail in trends:
    para = doc.add_paragraph()
    run = para.add_run(f"{title_text}: ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    detail_run = para.add_run(detail)
    detail_run.font.size = Pt(11)
    detail_run.font.name = 'Calibri'

# ── Prompt Improvement Suggestions ───────────────────────────────────────────

doc.add_heading('Prompt Improvement Suggestions', level=2)
for run in doc.paragraphs[-1].runs:
    run.font.size = Pt(13)
    run.font.bold = True

suggestions = [
    "Add general price ranges to the agent's knowledge base for common services (teeth whitening, cleaning, checkups) so it can provide ballpark figures before recommending a consultation. This addresses the gap identified in Call #6.",
    "Introduce an earlier callback/priority-rebooking offer in the de-escalation flow. Currently the agent waits for a natural pause; prompting it to offer a solution within the first 30 seconds of detecting frustration would improve resolution speed (Call #3).",
    "Add softer framing templates for declining requests (e.g., personal contact info). Instead of a direct 'I can't share that,' use empathy-first language: 'For privacy reasons, I'm not able to share personal numbers, but I'd love to help you connect through the clinic.' (Call #10).",
    "Implement post-call confirmation tracking for offered next steps (consult bookings, transfers). This will close the resolution tracking gap identified in Calls #6 and #10 and provide better data for future QA audits.",
    "Consider adding a brief after-hours FAQ capability so the agent can answer common questions (hours, location, emergency info) even when booking is the primary after-hours function (reinforces Call #5's strong handling).",
]

for i, suggestion in enumerate(suggestions, 1):
    para = doc.add_paragraph()
    run = para.add_run(f"{i}. ")
    run.bold = True
    run.font.size = Pt(11)
    detail_run = para.add_run(suggestion)
    detail_run.font.size = Pt(11)
    detail_run.font.name = 'Calibri'
    para.paragraph_format.space_after = Pt(6)

# ── Comparison to Previous Period ────────────────────────────────────────────

doc.add_heading('Comparison to Previous Period', level=2)
for run in doc.paragraphs[-1].runs:
    run.font.size = Pt(13)
    run.font.bold = True

para = doc.add_paragraph()
note_run = para.add_run("Note: ")
note_run.bold = True
note_run.font.size = Pt(11)
detail_run = para.add_run("No prior period QA data is available for comparison. This is the first QA audit for the Bright Smile Dental AI Receptionist. Future reports will include period-over-period trending for all dimensions.")
detail_run.font.size = Pt(11)
detail_run.font.name = 'Calibri'
detail_run.italic = True

# ── Footer ───────────────────────────────────────────────────────────────────

# Add footer with "Prepared by Inspra AI"
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("Prepared by Inspra AI")
    run.font.size = Pt(9)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(128, 128, 128)

# ── Save ─────────────────────────────────────────────────────────────────────

output_dir = r"d:\Imp Claude Codes\Inspra Agent Skills\agent-ops-workspace\iteration-1\eval-reports-docx\with_skill\outputs"
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, "BrightSmileDental_QAReport_Feb2026.docx")
doc.save(output_path)
print(f"QA Report saved to: {output_path}")
print(f"Overall Score: {overall_score}/100")
print(f"Calls Audited: {len(calls)}")
print(f"Flagged Calls: {len(flagged_calls)}")
